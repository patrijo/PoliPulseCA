# ===================================================================================================================
# shared_utils.py
# Shared helper functions for text processing, transformer scoring,
# and statistics in PoliPulseCA.
# Dependencies: torch, transformers, nltk, vaderSentiment, textblob, matplotlib, seaborn, scipy, pandas
# Author: P. Jost
# Date: May 2025
# License: MIT License (code), CC BY 4.0 (training data/tables/docs) https://creativecommons.org/licenses/by/4.0/
#          Transformer models: Apache 2.0; data: US public domain/open
# ===================================================================================================================

# ─────────────────────────────────────────────────────────────────────────────
# SHARED UTILITIES FOR TEXT & MODEL PROCESSING (SENTIMENT/POLITONE)
# ─────────────────────────────────────────────────────────────────────────────

import os, io, re, zipfile, math
import torch, torch.nn as nn, torch.nn.functional as F
import numpy as np, pandas as pd, nltk
from docx import Document
from datetime import datetime
from collections import Counter
from vaderSentiment.vaderSentiment import SentimentIntensityAnalyzer
from textblob import TextBlob
from transformers import (
    AutoTokenizer, AutoModelForSequenceClassification,
    logging as hfl
)
from transformers.modeling_outputs import SequenceClassifierOutput  # <-- FIXED IMPORT HERE
import matplotlib.pyplot as plt
import seaborn as sns
from scipy.stats import spearmanr, mannwhitneyu
from nltk.corpus import stopwords

# ────── Initialization: NLTK data ──────
nltk.download("punkt", quiet=True)
nltk.download("stopwords", quiet=True)
torch.classes.__path__ = []  # Streamlit fix

# ────── Model Path Discovery ──────
SENTIMENT_MODELS_DIR = "./sentiment_model"
POLITONE_MODELS_DIR  = "./politone_model"

def _discover_transformer_paths():
    paths = {}
    if os.path.isdir(SENTIMENT_MODELS_DIR):
        for name in sorted(os.listdir(SENTIMENT_MODELS_DIR)):
            if os.path.isdir(os.path.join(SENTIMENT_MODELS_DIR, name)):
                paths[name] = name
    return paths

TRANSFORMER_PATHS = _discover_transformer_paths()
_STOPWORDS = set(stopwords.words('english'))

# ────── File Validation (GLOBAL CHECK) ──────
def check_uploaded_filenames(uploaded, st_warn=True, context=""):
    invalid_files = []
    valid_files = []
    pattern = r"^\d{4}-\d{2}-\d{2}_.+\.(txt|docx)$"
    for file in uploaded:
        if file.name.endswith(".zip"):
            with zipfile.ZipFile(file, "r") as z:
                for zipname in z.namelist():
                    # Only check actual files (not folders in zip)
                    if zipname.endswith((".txt", ".docx")):
                        if not re.match(pattern, os.path.basename(zipname)):
                            invalid_files.append(zipname)
                        else:
                            valid_files.append(("zip", file, zipname))
                    else:
                        invalid_files.append(zipname)
        else:
            if not re.match(pattern, file.name):
                invalid_files.append(file.name)
            else:
                valid_files.append(("direct", file, None))
    if st_warn and invalid_files:
        try:
            import streamlit as st
            st.warning(
                f"{context} The following files are invalid (wrong extension or filename):\n"
                + "\n".join(f"- `{f}`" for f in invalid_files)
                + "\n\n**Filenames must be of the form `yyyy-mm-dd_speaker.txt` or `.docx`.**"
            )
        except Exception:
            pass  # Allow use outside of Streamlit for testing
    return valid_files, invalid_files

# ────── Model Loading ──────
def load_transformer_model(path):
    hfl.set_verbosity_error()
    tokenizer = AutoTokenizer.from_pretrained(path)
    model     = AutoModelForSequenceClassification.from_pretrained(path)
    return tokenizer, model

def politone_forward(model):
    model.config.id2label = {0: "-1", 1: "0", 2: "1"}
    model.config.label2id = {"-1": 0, "0": 1, "1": 2}
    model.pre_classifier = nn.Identity()
    model.dropout = nn.Identity()
    def noReLUforward(self, input_ids=None, attention_mask=None, labels=None, **kwargs):
        hb     = self.distilbert(input_ids=input_ids, attention_mask=attention_mask, **kwargs)
        cls    = hb.last_hidden_state[:, 0]
        logits = self.classifier(cls)
        loss   = F.cross_entropy(logits, labels) if labels is not None else None
        return SequenceClassifierOutput(  # <-- FIXED REFERENCE HERE
            loss=loss, logits=logits, hidden_states=None, attentions=None
        )
    model.forward = noReLUforward.__get__(model, model.__class__)
    model.eval()

# ────── Text Cleaning & Extraction ──────
def clean_text(text):
    return re.sub(r'\s+', ' ', text).strip() if isinstance(text, str) else ""

def split_sentences(text):
    """Split text into sentences using NLTK."""
    return nltk.sent_tokenize(text)

def parse_filename(name):
    base = os.path.splitext(name)[0]
    m = re.match(r"(\d{4}-\d{2}-\d{2})_(.+)", base)
    if not m:
        return None, None
    date = datetime.strptime(m.group(1), "%Y-%m-%d")
    speaker = m.group(2).strip()
    return date, speaker

def extract_text(file):
    if file.name.endswith(".txt"):
        file.seek(0)
        return file.read().decode("utf-8")
    elif file.name.endswith(".docx"):
        file.seek(0)
        from docx import Document
        doc = Document(file)
        return "\n".join([p.text for p in doc.paragraphs])
    return ""

def extract_texts_from_zip(zip_file):
    records = []
    with zipfile.ZipFile(zip_file, "r") as z:
        for name in z.namelist():
            if name.endswith((".txt", ".docx")):
                date, speaker = parse_filename(name)
                if date:
                    with z.open(name) as f:
                        if name.endswith(".txt"):
                            txt = f.read().decode("utf-8")
                        elif name.endswith(".docx"):
                            file_bytes = f.read()  # read all bytes
                            doc = Document(io.BytesIO(file_bytes))
                            txt = "\n".join([p.text for p in doc.paragraphs])
                        else:
                            txt = ""
                        records.append({"date": date, "speaker": speaker, "text": clean_text(txt)})
    return pd.DataFrame(records)

# ────── Simple Sentiment Scoring ──────
def vader_sentiment(text):
    return SentimentIntensityAnalyzer().polarity_scores(text)['compound']

def textblob_sentiment(text):
    return TextBlob(text).sentiment.polarity

def mark_sig(p):
    if p < 0.001: return '***'
    if p < 0.01:  return '**'
    if p < 0.05:  return '*'
    return ''

# ────── Transformer Sentence Scoring & Postprocessing ──────
def apply_repetition_penalty(
    sentence, score, threshold=1, penalty_per_word=0.5, max_penalty=1.0, exclude_stopwords=False
):
    tokens = re.findall(r'\b\w+\b', sentence.lower())
    if exclude_stopwords:
        tokens = [t for t in tokens if t not in _STOPWORDS]
    counts = Counter(tokens)
    repeated = sum(max(0, c - threshold) for c in counts.values())
    factor = 1 - min(repeated * penalty_per_word, max_penalty)
    return score * factor

def apply_winsorizing(scores, threshold=2.5):
    m, std = scores.mean(), scores.std()
    return np.clip(scores, m - threshold * std, m + threshold * std)

def score_transformer_sentences(
    sentences, tokenizer, model, transformer_choice,
    apply_temp=False, temp_value=1.0,
    apply_conf=False,
    apply_repeat=False, repeat_thresh=1,
    repeat_penalty_value=0.5, repeat_max=1.0,
    exclude_stopwords=False,
    use_tanh=False, logit_scale=1.0,
    apply_ma=False, ma_window=3,
    apply_wins=False, wins_threshold=2.5,
    device=torch.device("cpu")
):
    scores = []
    for s in sentences:
        enc = tokenizer(s, return_tensors="pt", truncation=True, padding=True)
        enc = {k: v.to(device) for k, v in enc.items()}
        with torch.no_grad():
            logits = model(**enc).logits.squeeze()
            if apply_temp:
                logits = logits / temp_value
            probs = F.softmax(logits, dim=-1)
        n = probs.size(0)
        if n == 2:
            neg, neu, pos = probs[0].item(), 0.0, probs[1].item()
        elif n == 3:
            neg, neu, pos = probs.tolist()
        else:
            raise ValueError(f"Labels={n} unsupported")
        sc = pos - neg
        if apply_repeat:
            sc = apply_repetition_penalty(
                s, sc, repeat_thresh, repeat_penalty_value, repeat_max, exclude_stopwords
            )
        if apply_conf:
            sc *= probs.max().item()
        if use_tanh:
            sc = math.tanh(sc / logit_scale)
        scores.append(sc)
    arr = np.array(scores) if scores else np.array([0.0])
    if apply_ma and len(arr) >= ma_window:
        arr = np.convolve(arr, np.ones(ma_window)/ma_window, mode='same')
    if apply_wins:
        arr = apply_winsorizing(arr, wins_threshold)
    return float(arr.mean())

# ────── Statistical Helpers (Effect Size, MWU, Trends) ──────
def compute_mw_effect_size(u_stat, n1, n2):
    m  = n1 * n2 / 2
    sd = math.sqrt(n1 * n2 * (n1 + n2 + 1) / 12)
    z  = (u_stat - m) / sd
    return abs(z) / math.sqrt(n1 + n2)

def mann_whitney_between_groups(
    df,
    value_column="Transformer_Sentiment",
    group_column="Group"
):
    a = df[df[group_column] == df[group_column].unique()[0]][value_column]
    b = df[df[group_column] == df[group_column].unique()[1]][value_column]
    u, p = mannwhitneyu(a, b, alternative="two-sided")
    u = max(u, len(a) * len(b) - u)
    r = compute_mw_effect_size(u, len(a), len(b))
    return {
        "Groups Compared": f"{df[group_column].unique()[0]} vs {df[group_column].unique()[1]}",
        "U-statistic": u,
        "p-value": p,
        "Effect Size (r)": f"{r:.3f}",
        "Significance": mark_sig(p)
    }

def visualize_general_trend(selected_df):
    import streamlit as st
    st.subheader("Sentiment Trends")
    fig, ax = plt.subplots()
    years_numeric = pd.to_numeric(selected_df["year"])
    min_year, max_year = int(years_numeric.min()), int(years_numeric.max())
    for year in range(min_year, max_year + 1):
        ax.axvline(x=year, color="gray", linestyle="--", alpha=0.2, zorder=0)
    palette = sns.color_palette("deep", 3)
    methods = ["VADER_Sentiment", "TextBlob_Sentiment", "Transformer_Sentiment"]
    spearman_stats = []
    for i, method in enumerate(methods):
        dfm = selected_df.dropna(subset=[method])
        yrs = pd.to_numeric(dfm["year"])
        vals = dfm[method].astype(float)
        sns.lineplot(x=yrs, y=vals, label=method, ax=ax, color=palette[i], zorder=2)
        r, p = spearmanr(yrs, vals)
        z = np.polyfit(yrs, vals, 1)
        ax.plot(yrs, np.poly1d(z)(yrs), linestyle=":", alpha=0.8, color=palette[i],
                label=f"{method} Trend (ρ={r:.2f})", zorder=1)
        spearman_stats.append({
            "Method": method,
            "Spearman ρ": f"{r:.3f}",
            "p-value": p,
            "Significance": mark_sig(p)
        })
    ax.set_ylabel("Sentiment Score (−1 to 1)")
    ax.set_xlabel("Year")
    ax.set_ylim(-1, 1)
    ax.set_title("Sentiment Trends by Year")
    span = max_year - min_year
    step = 1 if span <= 10 else (5 if span <= 20 else 10)
    xticks = np.arange(min_year, max_year + 1, step)
    ax.set_xticks(xticks)
    ax.set_xticklabels(xticks)
    ax.legend(loc="lower center", bbox_to_anchor=(0.5, -0.35), ncol=2)
    ax.grid(True)
    st.pyplot(fig)
    stats_df = pd.DataFrame(spearman_stats)
    return fig, stats_df

def plot_dual_trends(
    result_df,
    x_col: str,
    y1_col: str,
    y2_col: str,
    title: str,
    ylabel: str):
    fig, ax = plt.subplots(figsize=(14, 6))
    palette = sns.color_palette("deep", 2)
    x = result_df[x_col].values
    sns.lineplot(x=x, y=result_df[y1_col], marker="o", label=y1_col, ax=ax, color=palette[0], zorder=2)
    z1 = np.polyfit(x, result_df[y1_col], 1)
    ax.plot(x, np.poly1d(z1)(x), linestyle=":", alpha=0.7, color=palette[0], label=f"{y1_col} Trend", zorder=1)
    sns.lineplot(x=x, y=result_df[y2_col], marker="o", label=y2_col, ax=ax, color=palette[1], zorder=2)
    z2 = np.polyfit(x, result_df[y2_col], 1)
    ax.plot(x, np.poly1d(z2)(x), linestyle=":", alpha=0.7, color=palette[1], label=f"{y2_col} Trend", zorder=1)
    ax.set_xlabel("Speech Index per Speaker")
    ax.set_ylabel(ylabel)
    ax.set_title(title, fontsize=14)
    ax.set_xticks(x)
    ax.legend(loc="lower center", bbox_to_anchor=(0.5, -0.25), ncol=2)
    ax.grid(True)
    return fig
